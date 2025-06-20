import sys
import os
from openai import OpenAI
from docx import Document
from datetime import datetime

# Initialize the OpenAI client
client = OpenAI()
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')

def moderation(task_description):
    moderation = client.moderations.create(input=task_description.lower())
    return moderation.results[0].flagged


def guardian_ai_task(task_description):
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": "Your role is to assess whether the user's question is allowed or not."},
            {"role": "user", "content": task_description}
        ]
    )
    return response.choices[0].message.content


def perform_web_search(query):
    """Use OpenAI's web_search_preview tool to fetch search results."""
    response = client.responses.create(
        model="gpt-4.1",
        tools=[{"type": "web_search_preview"}],
        input=query
    )
    return response.output_text


#You are the Queen AI, overseeing and managing all other AIs."
    
def queen_ai_task(task_description):
    response = client.chat.completions.create(
        model="o4-mini",
        messages=[
            {"role": "user", "content": f"""
             You are the Queen AI, the central orchestrator of a network of specialized AI agents. Your primary purpose is to decompose complex user requests into manageable subtasks, delegate these subtasks to the most capable agents, guide them as necessary, synthesize their outputs, and deliver a coherent, complete, and optimal solution to the end user.
Core Objectives:Clarity and Comprehensiveness: Break down each user request into clear, logical subtasks. Ensure your directives are easy for subordinate agents to interpret and execute.

Efficiency and Quality Control: Assign subtasks to appropriate agents based on their known capabilities and expertise. Continuously monitor their progress, ensuring that all outputs meet high standards of accuracy, usefulness, fairness, and ethical conduct.
Iterative Refinement: If intermediate results are unclear or incomplete, ask agents targeted questions or reassign tasks with more precise instructions. Encourage iterative improvement until the final result is of exceptional quality.
Harmonization and Integration: Combine all partial results into a cohesive, well-structured final answer. Preserve logical consistency, ensure no contradictions, and highlight the most important insights.
Upholding Values: Always operate with care, respect, and truthfulness. Promote fairness, avoid biased or harmful content, and remain neutral, logical, and idealistic. Prioritize the user’s needs while maintaining high ethical and intellectual standards.

Decision-Making Criteria:

Relevance: Only produce directives and outputs that directly serve the user’s goals.
Accuracy: Rely on verified information and sound reasoning. In cases of uncertainty, clarify assumptions or request additional input from expert agents.
Elegance and Utility: Present final results in a polished, easy-to-understand format that maximizes the end user’s value.
Security and Compliance: Follow all established guidelines and policies, ensuring that no harmful, unethical, or illicit behavior occurs within the subordinate agents’ actions.

Operational Methodology:

Initial Analysis: Begin by restating the user’s request and dissect it into specific, actionable subtasks or questions.
Agent Delegation: Assign each subtask to the most suitable agent, providing detailed instructions and quality expectations.
Progress Review: Periodically request and review updates from each agent, asking follow-up questions to clarify ambiguities or correct shortcomings.
Synthesis and Validation: Integrate all agent outputs into a single, refined answer. Check for coherence, ethical soundness, correctness, and completeness.
Final Delivery: You will present the final, well-structured response to the user, ensuring it is both helpful and fully aligned with the user’s initial requirements.
Task: {task_description}"""}
        ]
    )
    return response.choices[0].message.content


def sub_subordinate_ai_task(task_name, context):
    response = client.chat.completions.create(
        model="o4-mini",
        messages=[
            # {"role": "developer", "content": f"You are a sub-subordinate AI executing the task: {task_name}."},
            {"role": "user", "content": f"You are a sub-subordinate AI executing the task: {task_name}. {context}"}
        ]
    )
    return response.choices[0].message.content


def subordinate_ai_task(task_name, queen_instruction):
    context = f"Original Queen AI instruction: {queen_instruction}\nSubordinate AI task: {task_name}"
    response = client.chat.completions.create(
        model="gpt-4.1",
        messages=[
            {"role": "developer", "content": f"You are a subordinate AI executing the following task: {task_name}."},
            {"role": "user", "content": context}
        ]
    )
    return response.choices[0].message.content


def computer_use_agent_task(task_name, context):
    """Invoke the OpenAI computer-use-preview model for browser-based tasks."""
    response = client.responses.create(
        model="computer-use-preview",
        tools=[{
            "type": "computer_use_preview",
            "display_width": 1024,
            "display_height": 768,
            "environment": "browser"
        }],
        input=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": f"{task_name}. {context}"}
                ]
            }
        ],
        reasoning={"summary": "concise"},
        truncation="auto"
    )
    return response.output


def dynamic_worker_assignment(worker_count, task_name, context, sub_agent_type=None):
    results = []
    for i in range(worker_count):
        response = client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": f"You are Worker {i+1} tasked with the following job: {task_name}."},
                {"role": "user", "content": context}
            ]
        )
        worker_result = response.choices[0].message.content
        if sub_agent_type:
            sub_context = f"Worker {i+1} result: {worker_result}. {context}"
            if sub_agent_type == "search":
                sub_response = perform_web_search(task_name)
                worker_result += f"\nWeb search agent result: {sub_response}"
            elif sub_agent_type == "computer":
                sub_response = computer_use_agent_task(task_name, sub_context)
                worker_result += f"\nComputer agent result: {sub_response}"
            else:  # generic helper
                sub_response = sub_subordinate_ai_task(task_name, sub_context)
                worker_result += f"\nSub-agent response: {sub_response}"
        results.append(worker_result)
    return results


def parse_task_description(task_description):
    """Check if the user wants to create multiple workers."""
    if 'send' in task_description.lower() and 'agents' in task_description.lower():
        parts = task_description.split(" ")
        for i, word in enumerate(parts):
            if word.isdigit():
                return int(word)
    return 1  # Default to one agent if not specified


def parse_worker_agent_type(task_description):
    """Determine what type of helper agent workers should use, if any."""
    lower = task_description.lower()
    if "computer use agent" in lower:
        return "computer"
    if "web search agent" in lower:
        return "search"
    if "workers can use agents" in lower:
        return "general"
    return None


def save_response(task_name, task_description, response, doc_name="HiveMinds_output.docx"):
    """Save the response either to a Word document or to a code file."""
    desc_lower = task_description.lower()
    if any(k in desc_lower for k in ["blog", "email", "research"]):
        document = Document(doc_name) if os.path.exists(doc_name) else Document()
        document.add_heading(task_name, level=1)
        document.add_paragraph(response)
        document.add_paragraph("Saved: " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        document.save(doc_name)
    elif "code" in desc_lower:
        file_name = f"{task_name.replace(' ', '_')}.py"
        with open(file_name, "w") as f:
            f.write(response)
    else:
        document = Document(doc_name) if os.path.exists(doc_name) else Document()
        document.add_heading(task_name, level=1)
        document.add_paragraph(response)
        document.add_paragraph("Saved: " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        document.save(doc_name)


def collect_tasks():
    tasks = {}
    print("Enter tasks for the Queen AI to delegate. Type 'done' when finished.")
    while True:
        task_name = input("Enter task name: ")
        if task_name.lower() == 'done':
            break
        task_description = input("Enter task description: ")
        tasks[task_name] = task_description
    return tasks


tasks = collect_tasks()
queen_responses = {}
subordinate_responses = {}
dynamic_worker_responses = {}
queen_final_responses = {}

for task_name, task_description in tasks.items():
    print("\nHiveMindsAI processing queen task:", task_name)
    worker_count = parse_task_description(task_description)
    sub_agent_type = parse_worker_agent_type(task_description)
    queen_instruction = queen_ai_task(f"Provide detailed instructions for a subordinate AI to: {task_description}")
    queen_responses[task_name] = queen_instruction
    print(f"\nQueen AI response for {task_name}:\n{queen_instruction}")

    if worker_count > 1 or sub_agent_type:
        print(f"\nAssigning {worker_count} agents to handle the task: {task_name}")
        if sub_agent_type:
            print(f"Workers will utilize {sub_agent_type} agents for assistance")
        context = f"Task: {task_description} as instructed by the Queen AI."
        worker_responses = dynamic_worker_assignment(worker_count, task_name, context, sub_agent_type=sub_agent_type)
        dynamic_worker_responses[task_name] = worker_responses
        combined = "\n".join(worker_responses)
        queen_final_responses[task_name] = queen_ai_task(
            f"Integrate the following results and produce a final explanation for '{task_description}':\n{combined}"
        )
        save_response(task_name, task_description, queen_final_responses[task_name])

    else:
        print("\nHiveMindsAI processing subordinate task:", task_name)
        subordinate_response = subordinate_ai_task(task_name, queen_instruction)
        subordinate_responses[task_name] = subordinate_response
        print(f"\nSubordinate AI response for {task_name}:\n{subordinate_response}")
        save_response(task_name, task_description, subordinate_response)

print("\nSummary of tasks and responses:")
for task_name in tasks.keys():
    print(f"\nTask: {task_name}")
    print(f"Queen AI response: {queen_responses[task_name]}")
    if task_name in dynamic_worker_responses:
        print(f"Dynamic Worker responses: {dynamic_worker_responses[task_name]}")
        print(f"Queen final response: {queen_final_responses[task_name]}")
    elif task_name in subordinate_responses:
        print(f"Subordinate AI response: {subordinate_responses[task_name]}")
